import os
import sys
import tkinter as tk
from tkinter import filedialog, messagebox, simpledialog
from pdf_extraction.extractor import extract_text_from_pdf
from text_analysis.txt_analysis import summarize_text, extract_keywords, categorize_text, extract_text_with_ocr
from docx import Document
from text_analysis.query_and_train_model import query_and_train_model
from word_generation.text_generation import train_model_on_corpus

# Add the src directory to Python's path
sys.path.append(os.path.abspath(os.path.join(os.path.dirname(__file__), "..")))

# Global variable to store extracted text and file path
extracted_text = ""
pdf_file_path = ""
CORPUS_DIR = "PDAI/data/corpus"
CORPUS_DIR = "PDAI/data/corpus"

# Function to handle file selection and text extraction
def select_file_and_extract():
    global extracted_text, pdf_file_path
    pdf_file_path = filedialog.askopenfilename(filetypes=[("PDF Files", "*.pdf")])
    
    if pdf_file_path:
        extracted_text_dict = extract_text_from_pdf(pdf_file_path)
        if extracted_text_dict:
            messagebox.showinfo("Extraction Success", "Text extracted successfully!")
            text_output.delete(1.0, tk.END)
            
            # Join all pages' text into a single string for further analysis
            extracted_text = ""
            for page_num, page_text in extracted_text_dict.items():
                extracted_text += f"--- Page {page_num} ---\n{page_text}\n\n"
                text_output.insert(tk.END, f"--- Page {page_num} ---\n{page_text}\n\n")
            
            # Clear intermediate memory
            del extracted_text_dict
        else:
            messagebox.showerror("Extraction Failed", "Unable to extract text from the PDF.")
    else:
        messagebox.showwarning("No File Selected", "Please select a PDF file to extract.")

# Function to export the extracted text to a .txt file
# Function to export the extracted text to a .txt file
def export_to_txt():
    global extracted_text, pdf_file_path
    if extracted_text and pdf_file_path:
        file_name = os.path.splitext(os.path.basename(pdf_file_path))[0]
        txt_file_path = os.path.join(os.path.dirname(pdf_file_path), f"{file_name}.txt")
        try:
            with open(txt_file_path, 'w', encoding='utf-8') as txt_file:
                txt_file.write(extracted_text)
            messagebox.showinfo("Export Success", f"Text exported successfully to {txt_file_path}")
        except Exception as e:
            messagebox.showerror("Export Failed", f"Failed to export text: {e}")
    else:
        messagebox.showwarning("No Text to Export", "Please extract text from a PDF first.")

# Function to export the extracted text to a .docx file
# Function to export the extracted text to a .docx file
def export_to_docx():
    global extracted_text, pdf_file_path
    if extracted_text and pdf_file_path:
        file_name = os.path.splitext(os.path.basename(pdf_file_path))[0]
        docx_file_path = os.path.join(os.path.dirname(pdf_file_path), f"{file_name}.docx")
        try:
            doc = Document()
            doc.add_heading('Extracted Data', level=1)
            for page in extracted_text.split("\n\n"):
                doc.add_paragraph(page)
            doc.save(docx_file_path)
            messagebox.showinfo("Export Success", f"Text exported successfully to {docx_file_path}")
        except Exception as e:
            messagebox.showerror("Export Failed", f"Failed to export text: {e}")
    else:
        messagebox.showwarning("No Text to Export", "Please extract text from a PDF first.")

# NLP Functions to analyze extracted text
def summarize_extracted_text():
    global extracted_text
    if extracted_text:
        summary = summarize_text(extracted_text)
        text_output.insert(tk.END, f"\n--- Summary ---\n{summary}\n")
    else:
        messagebox.showwarning("No Text to Summarize", "Please extract text from a PDF first.")

def extract_keywords_from_text():
    global extracted_text
    if extracted_text:
        keywords = extract_keywords(extracted_text)
        text_output.insert(tk.END, f"\n--- Keywords ---\n{', '.join(keywords)}\n")
    else:
        messagebox.showwarning("No Text for Keywords", "Please extract text from a PDF first.")

def categorize_extracted_text():
    global extracted_text
    if extracted_text:
        categories = categorize_text(extracted_text)
        text_output.insert(tk.END, f"\n--- Categories ---\n{', '.join(categories)}\n")
    else:
        messagebox.showwarning("No Text to Categorize", "Please extract text from a PDF first.")

# OCR Button to handle OCR extraction
def perform_ocr_extraction():
    global pdf_file_path
    if pdf_file_path:
        ocr_text = extract_text_with_ocr(pdf_file_path)
        text_output.insert(tk.END, f"\n--- OCR Text ---\n{ocr_text}\n")
    else:
        messagebox.showwarning("No File Selected", "Please select a PDF file for OCR.")

# Create the main window
root = tk.Tk()
root.title("PDF Text Extractor with NLP, OCR Tools and Corpus Training")

# Create buttons for actions
select_file_button = tk.Button(root, text="Select PDF", command=select_file_and_extract)
summarize_button = tk.Button(root, text="Summarize Text", command=summarize_extracted_text)
keywords_button = tk.Button(root, text="Extract Keywords", command=extract_keywords_from_text)
categorize_button = tk.Button(root, text="Categorize Text", command=categorize_extracted_text)
ocr_button = tk.Button(root, text="Perform OCR", command=perform_ocr_extraction)
export_to_docx_button = tk.Button(root, text="Export to .docx", command=export_to_docx)
export_to_txt_button = tk.Button(root, text="Export to .txt", command=export_to_txt)
train_button = tk.Button(root, text="Train Model on Corpus", command=query_and_train_model)

# Place buttons in a single row (horizontal alignment) at the top using grid
buttons = [
    select_file_button, summarize_button, keywords_button, categorize_button,
    ocr_button, export_to_docx_button, export_to_txt_button, train_button
]

for col_index, button in enumerate(buttons):
    button.grid(row=0, column=col_index, padx=5, pady=5, sticky="ew")

# Configure the grid columns to adjust equally with window resizing
for col_index in range(len(buttons)):
    root.grid_columnconfigure(col_index, weight=1)

# Text widget to display extracted text, separated by pages
text_output = tk.Text(root, height=20, width=80)
text_output.grid(row=1, column=0, columnspan=8, padx=5, pady=5, sticky="nsew")

# Configure the row for resizing
root.grid_rowconfigure(1, weight=1)

# Run the Tkinter event loop
root.mainloop()

