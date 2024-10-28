from docx import Document
from docx.shared import Inches

def generate_word_doc(text_data, images, output_path="output.docx"):
    doc = Document()
    doc.add_heading('Extracted Data', level=1)

    # Add text to the document
    for page_num, page_text in text_data.items():
        doc.add_heading(f"Page {page_num}", level=2)
        doc.add_paragraph(page_text)

    # Add images to the document
    for img_path in images:
        doc.add_picture(img_path, width=Inches(4.0))  # Adjust size as needed

    doc.save(output_path)
